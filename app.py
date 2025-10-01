import streamlit as st
import google.generativeai as genai
from PIL import Image
import pandas as pd
import json
import os
from dotenv import load_dotenv
import docx
import fitz  # PyMuPDF
import io
import re

# Load environment variables from a .env file
load_dotenv()

def configure_api():
    """
    Configures the Google Generative AI API with the key from environment variables.
    Handles errors if the API key is not found.
    """
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        st.error("GEMINI_API_KEY environment variable not set. Please create a .env file and add your API key.")
        st.stop()
    genai.configure(api_key=api_key)

def list_available_models():
    """
    Lists available Gemini models that support 'generateContent' and displays them in a table.
    """
    try:
        st.info("Fetching available models...")
        models_list = []
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                models_list.append({'Model Name': m.name, 'Description': m.description})
        
        if models_list:
            df = pd.DataFrame(models_list)
            st.dataframe(df, use_container_width=True)
        else:
            st.warning("No models supporting 'generateContent' were found for your API key.")
    except Exception as e:
        st.error(f"Could not retrieve the model list: {e}")

def get_gemini_response(image, prompt):
    """
    Calls the Gemini API to get text and table data from an image.
    """
    model = genai.GenerativeModel('gemini-2.5-flash') #gemini-1.5-flash-latest
    try:
        response = model.generate_content([prompt, image])
        return response.text
    except Exception as e:
        st.error(f"An error occurred with the Gemini API: {e}")
        return None

def extract_data_from_docx(file_stream, prompt):
    """
    Extracts text, native tables, and data from embedded images in a .docx file.
    Each table is stored as a separate dictionary in a list.
    """
    doc = docx.Document(file_stream)
    full_text = [para.text for para in doc.paragraphs]
    all_tables_data = []

    # 1. Extract native Word tables
    for i, table in enumerate(doc.tables):
        table_data = [[cell.text for cell in row.cells] for row in table.rows]
        all_tables_data.append({'title': f'Native Table {i+1}', 'data': table_data})

    # 2. Find and extract tables from embedded images
    image_parts = [
        rel.target_part
        for rel in doc.part.rels.values()
        if "image" in rel.target_ref
    ]

    if image_parts:
        status = st.status(f"Found {len(image_parts)} embedded images. Analyzing...", expanded=True)
        for i, part in enumerate(image_parts):
            status.write(f"Analyzing embedded image {i + 1}...")
            try:
                image = Image.open(io.BytesIO(part.blob))
                response_text = get_gemini_response(image, prompt)
                
                if response_text:
                    start_index = response_text.find('{')
                    end_index = response_text.rfind('}')
                    if start_index != -1 and end_index != -1:
                        json_string = response_text[start_index:end_index+1]
                        data = json.loads(json_string)
                        img_text = data.get("text", "")
                        img_table = data.get("table", [])

                        if img_text:
                            full_text.append(f"\n--- Text from Embedded Image {i + 1} ---\n{img_text}")
                        if img_table:
                            all_tables_data.append({'title': f'Table from Embedded Image {i+1}', 'data': img_table})
            except Exception as e:
                st.warning(f"Could not process embedded image {i + 1}. It might be a non-standard format. Error: {e}")
        status.update(label="Image analysis complete!", state="complete")

    return "\n".join(full_text), all_tables_data


def extract_data_from_pdf(file_stream, prompt):
    """
    Extracts text and tables from a PDF by converting pages to images.
    Each table is stored as a separate dictionary in a list.
    """
    pdf_document = fitz.open(stream=file_stream.read(), filetype="pdf")
    aggregated_text = ""
    aggregated_tables = []

    status = st.status(f"Processing {len(pdf_document)} pages...", expanded=True)

    for page_num, page in enumerate(pdf_document):
        status.write(f"Analyzing page {page_num + 1}...")
        
        pix = page.get_pixmap()
        image = Image.open(io.BytesIO(pix.tobytes("png")))
        
        response_text = get_gemini_response(image, prompt)
        
        if response_text:
            try:
                start_index = response_text.find('{')
                end_index = response_text.rfind('}')
                if start_index != -1 and end_index != -1:
                    json_string = response_text[start_index:end_index+1]
                    data = json.loads(json_string)
                    
                    page_text = data.get("text", "")
                    page_table = data.get("table", [])
                    
                    if page_text:
                        aggregated_text += f"\n\n--- Page {page_num + 1} ---\n{page_text}"
                    if page_table:
                        aggregated_tables.append({'title': f'Table from Page {page_num + 1}', 'data': page_table})
            except json.JSONDecodeError:
                st.warning(f"Could not parse data from page {page_num + 1}. Skipping.")

    status.update(label="PDF processing complete!", state="complete")
    return aggregated_text, aggregated_tables


def main():
    st.set_page_config(page_title="File Content Extractor", layout="wide", page_icon="📄")

    st.title("📄 File Content & Table Extractor")
    st.markdown("Upload an Image, PDF, or Word document to extract its text and structured table data.")

    with st.expander("Advanced Options"):
        if st.button("List Available Models"):
            configure_api()
            list_available_models()

    st.markdown("---")
    uploaded_file = st.file_uploader("Choose a file...", type=["jpg", "jpeg", "png", "pdf", "docx"])

    if uploaded_file is not None:
        col1, col2 = st.columns(2)
        
        with col1:
            if uploaded_file.type.startswith('image/'):
                st.image(Image.open(uploaded_file), caption="Uploaded Image", use_column_width=True)
            else:
                st.info(f"📄 Uploaded file: **{uploaded_file.name}**")
                st.markdown("Preview is not available. Click 'Extract Data' to process.")

        with col2:
            if st.button("✨ Extract Data", use_container_width=True):
                with st.spinner("Analyzing file..."):
                    configure_api()
                    
                    file_extension = os.path.splitext(uploaded_file.name)[1].lower()
                    text_result, table_result = "", []
                    prompt_template = "Extract all text and the primary table from this image as a JSON object with 'text' and 'table' keys. The 'table' value should be a list of lists."

                    if file_extension in [".jpg", ".jpeg", ".png"]:
                        response_text = get_gemini_response(Image.open(uploaded_file), prompt_template)
                        if response_text:
                            try:
                                start_index = response_text.find('{'); end_index = response_text.rfind('}')
                                if start_index != -1 and end_index != -1:
                                    data = json.loads(response_text[start_index:end_index+1])
                                    text_result = data.get("text", "")
                                    # Wrap single table in the expected list-of-dicts format
                                    if data.get("table"):
                                        table_result = [{'title': 'Table from Image', 'data': data.get("table")}]
                            except json.JSONDecodeError: st.error("Failed to decode JSON from API.")
                    
                    elif file_extension == ".pdf":
                        text_result, table_result = extract_data_from_pdf(uploaded_file, prompt_template)

                    elif file_extension == ".docx":
                        prompt_for_embeds = "This image was embedded in a document. Analyze it for tables. Provide output as a JSON object with 'text' and 'table' keys. The 'table' should be a list of lists."
                        text_result, table_result = extract_data_from_docx(uploaded_file, prompt_for_embeds)
                    
                    st.session_state.extracted_text = text_result
                    st.session_state.extracted_tables = table_result # Note the plural 'tables'
                    st.success("Data extracted successfully!")

    if 'extracted_text' in st.session_state:
        st.markdown("---"); st.header("Extracted Results")
        st.subheader("📝 Extracted Text")
        st.text_area("Text Content", st.session_state.extracted_text, height=250)
        
        tables_data = st.session_state.get('extracted_tables', [])
        if tables_data:
            st.subheader("📊 Extracted Tables")
            for table_info in tables_data:
                with st.container(border=True):
                    st.markdown(f"**{table_info['title']}**")
                    try:
                        table_rows = table_info['data']
                        if table_rows and len(table_rows) > 1:
                            df = pd.DataFrame(table_rows[1:], columns=table_rows[0])
                            st.dataframe(df)
                        elif table_rows:
                             st.dataframe(pd.DataFrame(table_rows))
                        else:
                            st.info("Table is empty.")
                    except Exception:
                        st.warning("Could not format table. Displaying raw data.")
                        st.json(table_info['data'])
        else:
            st.info("No tables were found in the file.")

        st.markdown("---"); st.subheader("⬇️ Download Data")
        
        # Build the full text content for the TXT download
        txt_content = f"Extracted Text\n{'='*20}\n{st.session_state.extracted_text}"
        if tables_data:
            for table_info in tables_data:
                txt_content += f"\n\n\n{table_info['title']}\n{'='*20}\n"
                txt_content += pd.DataFrame(table_info['data']).to_string(index=False, header=False)
        
        st.download_button("Download All as TXT", txt_content, "extracted_content.txt", "text/plain", use_container_width=True)
        
        if tables_data:
            st.markdown("---")
            st.subheader("⬇️ Download Individual Tables as CSV")
            for i, table_info in enumerate(tables_data):
                try:
                    table_rows = table_info['data']
                    if table_rows and len(table_rows) > 1:
                        # Sanitize title for filename
                        safe_filename = re.sub(r'[^a-z0-9_]+', '', table_info['title'].lower().replace(' ', '_'))
                        csv = pd.DataFrame(table_rows[1:], columns=table_rows[0]).to_csv(index=False).encode('utf-8')
                        st.download_button(
                            label=f"Download '{table_info['title']}'",
                            data=csv,
                            file_name=f"{safe_filename}.csv",
                            mime="text/csv",
                            key=f"csv_dl_{i}" # Unique key is important for multiple buttons
                        )
                except Exception:
                    st.warning(f"Could not generate CSV for '{table_info['title']}'.", icon="⚠️")

if __name__ == "__main__":
    main()

